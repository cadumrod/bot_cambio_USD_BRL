from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support import expected_conditions
from selenium.common.exceptions import *
from selenium.webdriver.support.ui import WebDriverWait
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt, Cm
from hyperlink import add_hyperlink
from time import sleep
from docx2pdf import convert
import os
import platform
import ctypes
from ctypes import wintypes


##################################  Config Selenium  ##################################
def iniciar_driver():
    chrome_options = Options()
    arguments = ['--lang=pt-BR', '--window-size=1920,1080',
                 '--incognito', '--headless']
    for argument in arguments:
        chrome_options.add_argument(argument)

    chrome_options.add_experimental_option('prefs', {
        'download.prompt_for_download': False,
        'profile.default_content_setting_values.notifications': 2,
        'profile.default_content_setting_values.automatic_downloads': 1,

    })
    driver = webdriver.Chrome(service=ChromeService(
        ChromeDriverManager().install()), options=chrome_options)

    wait = WebDriverWait(
        driver,
        10,
        poll_frequency=1,
        ignored_exceptions=[
            NoSuchElementException,
            ElementNotVisibleException,
            ElementNotSelectableException
        ]
    )

    return driver, wait


##################################  Padronizando path da Área de trabalho  ##################################

# Definir o diretório da área de trabalho do usuário de forma compatível com diferentes sistemas operacionais para output de arquivos
def get_desktop_path():
    if platform.system() == "Windows":
        # Windows: Usar SHGetFolderPathW para encontrar a área de trabalho
        CSIDL_DESKTOP = 0x0000       # Pasta da área de trabalho
        SHGFP_TYPE_CURRENT = 0       # Pegar atual, não valor padrão

        buf = ctypes.create_unicode_buffer(wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(
            None, CSIDL_DESKTOP, None, SHGFP_TYPE_CURRENT, buf)
        return buf.value

    elif platform.system() == "Darwin":
        # macOS: Caminho padrão para a área de trabalho
        return os.path.join(os.path.expanduser("~"), "Desktop")

    elif platform.system() == "Linux":
        # Linux: Caminho padrão para a área de trabalho
        return os.path.join(os.path.expanduser("~"), "Área de Trabalho")

    else:
        raise Exception("Sistema Operacional não suportado.")


# Obter o caminho da área de trabalho
desktop_dir = get_desktop_path()


##################################  Execução do app  ##################################
print('#'*20, 'Monitoramento de câmbio', '#'*20)
sleep(1.5)
print("Iniciando monitoramento...")
driver, wait = iniciar_driver()
driver.get('https://br.investing.com/currencies/usd-brl')


# Aguardar retorno do container cotação para prosseguir
wait.until(expected_conditions.visibility_of_all_elements_located(
    (By.XPATH, "//div[@class='text-5xl/9 font-bold text-[#232526] md:text-[42px] md:leading-[60px]']")))

# Obter o valor da cotacao
elemento_cotacao = driver.find_element(
    By.XPATH, "//div[@class='text-5xl/9 font-bold text-[#232526] md:text-[42px] md:leading-[60px]']")

# Transformando elemento da cotacao em texto
texto_cotacao = elemento_cotacao.text

# Formatando para conversao em float e ajuste de casas decimais
texto_cotacao_limpo = texto_cotacao.replace(',', '.').strip()
float_cotacao = float(texto_cotacao_limpo)
cotacao_arredondada = round(float_cotacao, 2)

# Formatando a troca de ponto por vírgula
cotacao_formatada = f"{cotacao_arredondada:.2f}".replace('.', ',')

# Salvando data atual
data = datetime.strftime(datetime.now(), "%d/%m/%Y")

# Link de onde foi extraída a cotação
link_cotacao = driver.current_url

# Caminho completo para salvar o arquivo .png
output_file = os.path.join(desktop_dir, 'cotacao.png')

# Print em resolução 1920 x 1080
driver.save_screenshot(output_file)

# Encerrar conexão
driver.quit()

print("\nDados obtidos! Criando arquivo em formato .docx...")
sleep(1.5)


##################################  Criação do .docx  ##################################
# Criar documento docx
doc = Document()

# Título
titulo = doc.add_heading(level=1)
run = titulo.add_run(
    f'Cotação atual do Dólar - R${cotacao_formatada} ({data})')
run.bold = True
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(20)
titulo.paragraph_format.space_after = Pt(12)


# Informações em parágrafos
# Valor e data
paragrafo_texto = doc.add_paragraph()
run_paragrafo_1 = paragrafo_texto.add_run("O dólar está no valor de R$")
run_paragrafo_1.font.size = Pt(12)
run_paragrafo_1.font.name = 'Liberation Serif'

run_paragrafo_2 = paragrafo_texto.add_run(cotacao_formatada)
run_paragrafo_2.bold = True
run_paragrafo_2.font.size = Pt(12)
run_paragrafo_2.font.name = 'Liberation Serif'

run_paragrafo_3 = paragrafo_texto.add_run(f" na data {data}. \n")
run_paragrafo_3.font.size = Pt(12)
run_paragrafo_3.font.name = 'Liberation Serif'

# Link do site
run_paragrafo_4 = paragrafo_texto.add_run(
    f"Valor cotado no site: ")
add_hyperlink(paragrafo_texto, link_cotacao)
paragrafo_texto.add_run().add_break()
run_paragrafo_4.font.size = Pt(12)
run_paragrafo_4.font.name = 'Liberation Serif'
paragrafo_texto.paragraph_format.space_after = Pt(12)

# Texto de print da cotação
run_paragrafo_5 = paragrafo_texto.add_run(
    f"Segue o print da cotação atual:\n")
run_paragrafo_5.font.size = Pt(12)
run_paragrafo_5.font.name = 'Liberation Serif'

# Adicionando print como imagem
paragrafo_imagem = doc.add_paragraph()
run_imagem = paragrafo_imagem.add_run()
run_imagem.add_picture(output_file, width=Cm(15))
paragrafo_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Autor da cotacao
paragrafo_autor = doc.add_paragraph()
run_autor1 = paragrafo_autor.add_run(
    "Cotação realizada por - ")
run_autor1.font.size = Pt(12)
run_autor1.font.name = 'Liberation Serif'

run_autor2 = paragrafo_autor.add_run(
    "Carlos Rodrigues")
run_autor2.font.size = Pt(12)
run_autor2.font.name = 'Liberation Serif'
run_autor2.italic = True
run_autor2.bold = True

# Salvar o documento
doc.save(os.path.join(desktop_dir, 'cotacao.docx'))

# Encerrando criação de docx
sleep(1)
print("Arquivo docx pronto!")
sleep(1)


##################################  Conversão de .docx para .pdf  ##################################
print("Realizando conversão para .pdf...")
sleep(1)

convert(os.path.join(desktop_dir, 'cotacao.docx'))
sleep(1)

print("Conversão realizada com sucesso!")
sleep(1)

##################################  Encerramento do app  ##################################
print("Os arquivos foram salvos na sua Área de trabalho!")
sleep(1.5)
print("Encerrando em...")
for i in range(3, 0, -1):
    print(i)
    sleep(1)
os._exit(0)
